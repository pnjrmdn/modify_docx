import os
from docx import Document
from docx.enum.text import WD_UNDERLINE

def is_bold_and_underline_run(run):
    return run.bold and (run.underline == WD_UNDERLINE.SINGLE)

def process_docx_file(file_path):
    document = Document(file_path)
    
    for paragraph in document.paragraphs:
        if "SDN 017 BALIKPAPAN TIMUR" in paragraph.text:
            new_runs = []
            for run in paragraph.runs:
                if "SDN 017 BALIKPAPAN TIMUR" in run.text:
                    index = run.text.find("SDN 017 BALIKPAPAN TIMUR")
                    if index > 0:
                        prev_text = run.text[:index]
                        new_run = paragraph.add_run(prev_text)
                        new_run.bold = is_bold_and_underline_run(run)
                        new_run.underline = WD_UNDERLINE.SINGLE
                        new_runs.append(new_run)

                    new_run = paragraph.add_run("SDN 017 BALIKPAPAN TIMUR")
                    new_run.bold = True
                    new_run.underline = WD_UNDERLINE.SINGLE
                    new_runs.append(new_run)

                    remaining_text = run.text[index + len("SDN 017 BALIKPAPAN TIMUR"):]
                    if remaining_text:
                        new_run = paragraph.add_run(remaining_text)
                        new_run.bold = is_bold_and_underline_run(run)
                        new_run.underline = WD_UNDERLINE.SINGLE
                        new_runs.append(new_run)

                    run.text = ""

            for run in new_runs:
                paragraph.runs.append(run)
    
    document.save(file_path)

def main():
    current_directory = os.getcwd()
    
    for filename in os.listdir(current_directory):
        if filename.endswith(".docx"):
            file_path = os.path.join(current_directory, filename)
            process_docx_file(file_path)
            print(f"Processed: {filename}")

if __name__ == "__main__":
    main()
# test