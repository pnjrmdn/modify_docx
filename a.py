import os
import glob
from docx.shared import Pt
from docx import Document

# Define a function to set font style
def set_font(run):
    run.font.name = 'Arial'
    run.font.size = Pt(11)  # Set font size to 11 points

def replace_text_in_docx(file_path, old_text, new_text):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text, 1)  # Added "1" to limit replacements
            for run in paragraph.runs:
                set_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text, 1)  # Added "1" to limit replacements
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_font(run)
    doc.save(file_path)

directory = os.path.dirname(os.path.abspath(__file__))
old_text = "SD/MI  ( datadikdasmen.com )"
new_text = " SDN 017 BALIKPAPAN TIMUR"

del_text = "( RPP )"
del_text1 = "KURIKULUM 2013 (3 KOMPONEN) REVISI 2020 "
del_text2 = "(Sesuai Edaran Mendikbud Nomor 14 Tahun 2019) "
null_text = ""
docx_files = glob.glob(os.path.join(directory, "*.docx"))
for docx_file in docx_files:
    replace_text_in_docx(docx_file, del_text, null_text)
    replace_text_in_docx(docx_file, del_text1, null_text)
    replace_text_in_docx(docx_file, del_text2, null_text)
    replace_text_in_docx(docx_file, old_text, new_text)
    print(f"Text replaced and font set in {docx_file}")

#