import os
import glob
from docx.shared import Pt
from docx import Document

def replace_text_in_docx(file_path, old_text, new_text):
    doc = Document(file_path)
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text, 1)  # Added "1" to limit replacements
            for run in paragraph.runs:
                run.font.size = Pt(11)  # Set font size to 11 points
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text, 1)  # Added "1" to limit replacements
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)  # Set font size to 11 points
    doc.save(file_path)

directory = os.path.dirname(os.path.abspath(__file__))

old_text = "SD/MI  ( datadikdasmen.com )"
new_text = "SDN 017 BALIKPAPAN TIMUR"

old_kepsek = "SITI WARYATI,M.Pd."
new_kepsek = "AMINAH LINAWATI,S. Pd. MM"
old_teacher = "RIA PUTRI RAHMAWATI,M.Pd."
new_teacher = "SUNARTI IDA RISNAWATI, S. Pd. SD"
old_nip_kepsek = "NIP. 19720307 199603 2002"
new_nip_kepsek = "NIP 196906051999072001"
old_nip_teacher = "NIP. 19881014 201101 2008"
new_nip_teacher = "NIP 196807011991082002"

old_date = "Magelang, Juli 2020"
new_date = "Balikpapan, ………………. 20…"
del_text = "( RPP )"
del_text1 = "KURIKULUM 2013 (3 KOMPONEN) REVISI 2020 "
del_text2 = "(Sesuai Edaran Mendikbud Nomor 14 Tahun 2019) "
null_text = ""

docx_files = glob.glob(os.path.join(directory, "*.docx"))
for docx_file in docx_files:
    replace_text_in_docx(docx_file, del_text, null_text)
    replace_text_in_docx(docx_file, del_text1, null_text)
    replace_text_in_docx(docx_file, del_text2, null_text)
    replace_text_in_docx(docx_file, old_text, new_text)
    replace_text_in_docx(docx_file, old_kepsek, new_kepsek)
    replace_text_in_docx(docx_file, old_teacher, new_teacher)
    replace_text_in_docx(docx_file, old_nip_kepsek, new_nip_kepsek)
    replace_text_in_docx(docx_file, old_nip_teacher, new_nip_teacher)
    replace_text_in_docx(docx_file, old_date, new_date)
    print(f"Text replaced in {docx_file}")
